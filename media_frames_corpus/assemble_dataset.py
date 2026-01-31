"""
Assemble the Media Frames Corpus dataset from downloaded DOCX files.

Usage:
    python assemble_dataset.py [topic]

    topic: immigration (default), smoking, or samesex

Expected structure:
    media_frames_corpus/
        downloads/              <- for immigration
        smoking_downloads/      <- for smoking
        samesex_downloads/      <- for samesex
        immigration.json
        smoking.json
        samesex.json

Output:
    {topic}_corpus.parquet
    {topic}_corpus.csv
    {topic}_assembly_report.json
"""

import json
import sys
import os
import re
from pathlib import Path
from collections import defaultdict

import pandas as pd
from docx import Document

# Config - set by command line arg
CODES_JSON = "codes.json"


def get_topic_config(topic):
    """Get file paths for a given topic."""
    if topic == "immigration":
        downloads_dir = "downloads"
    else:
        downloads_dir = f"{topic}_downloads"

    return {
        "input_json": f"{topic}.json",
        "downloads_dir": downloads_dir,
        "output_parquet": f"{topic}_corpus.parquet",
        "output_csv": f"{topic}_corpus.csv",
        "report_file": f"{topic}_assembly_report.json",
    }


def normalize_title(title):
    """Normalize title for matching: lowercase, remove punctuation, collapse whitespace."""
    title = title.lower()
    title = re.sub(r'[^\w\s]', '', title)
    title = re.sub(r'\s+', ' ', title).strip()
    return title


def extract_text_from_docx(docx_path):
    """Extract body text from DOCX, skipping metadata header."""
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs]

    # Find "Body" marker and extract text after it
    body_idx = None
    for i, para in enumerate(paragraphs):
        if para.lower() == "body":
            body_idx = i
            break

    if body_idx is not None:
        body_text = "\n".join(paragraphs[body_idx + 1:])
    else:
        # Fallback: skip first ~10 lines (metadata) and take the rest
        body_text = "\n".join(paragraphs[10:])

    # Clean up
    body_text = body_text.strip()
    return body_text


def extract_title_from_docx(docx_path):
    """Extract the article title from DOCX content (usually 3rd non-empty paragraph)."""
    doc = Document(docx_path)
    non_empty = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # Title is typically the first substantial text after empty lines
    if len(non_empty) >= 1:
        return non_empty[0]
    return None


def load_nyt_articles(json_path):
    """Load NYT articles (excluding blogs) and build title lookup."""
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    nyt_articles = {
        k: v for k, v in data.items()
        if v.get('source', '').lower() == 'new york times'
    }

    # Build normalized title -> article mapping
    title_lookup = {}
    for article_id, article in nyt_articles.items():
        norm_title = normalize_title(article.get('title', ''))
        if norm_title in title_lookup:
            # Duplicate title - store as list
            if isinstance(title_lookup[norm_title], list):
                title_lookup[norm_title].append((article_id, article))
            else:
                title_lookup[norm_title] = [title_lookup[norm_title], (article_id, article)]
        else:
            title_lookup[norm_title] = (article_id, article)

    return nyt_articles, title_lookup


def load_codes(codes_path):
    """Load frame code mappings."""
    with open(codes_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def extract_frame_labels(article):
    """Extract frame codes per annotator from article annotations."""
    framing = article.get('annotations', {}).get('framing', {})

    annotator_frames = {}
    for annotator_key, annotations in framing.items():
        # Extract annotator ID (e.g., "annotator3_54.0_r" -> "annotator3")
        annotator_id = annotator_key.split('_')[0]

        # Collect unique base frame codes (1-15)
        frame_codes = set()
        for ann in annotations:
            code = ann.get('code')
            if code:
                base_code = int(float(code))  # 10.2 -> 10
                if 1 <= base_code <= 15:
                    frame_codes.add(base_code)

        annotator_frames[annotator_id] = sorted(frame_codes)

    return annotator_frames


def extract_tone_labels(article):
    """Extract tone labels per annotator."""
    tone = article.get('annotations', {}).get('tone', {})

    annotator_tones = {}
    for annotator_key, annotations in tone.items():
        annotator_id = annotator_key.split('_')[0]

        # Map tone codes
        tones = set()
        for ann in annotations:
            code = ann.get('code')
            if code:
                if 17.0 <= code < 18.0:
                    tones.add('pro')
                elif 18.0 <= code < 19.0:
                    tones.add('neutral')
                elif 19.0 <= code < 20.0:
                    tones.add('anti')

        annotator_tones[annotator_id] = list(tones)

    return annotator_tones


def main():
    # Parse topic from command line
    topic = sys.argv[1] if len(sys.argv) > 1 else "immigration"
    if topic not in ["immigration", "smoking", "samesex"]:
        print(f"Unknown topic: {topic}")
        print("Usage: python assemble_dataset.py [immigration|smoking|samesex]")
        sys.exit(1)

    config = get_topic_config(topic)
    print(f"Assembling {topic} corpus...")

    base_path = Path(__file__).parent
    os.chdir(base_path)

    # Load source data
    print("Loading source data...")
    nyt_articles, title_lookup = load_nyt_articles(config["input_json"])
    codes = load_codes(CODES_JSON)
    print(f"  Loaded {len(nyt_articles)} NYT articles")

    # Scan downloads directory
    downloads_path = base_path / config["downloads_dir"]
    if not downloads_path.exists():
        print(f"ERROR: Downloads directory not found: {downloads_path}")
        print(f"Please create the {config['downloads_dir']}/ folder and add batch subfolders with DOCX files.")
        return

    # Collect all DOCX files (case-insensitive, deduplicated)
    docx_files = []
    seen_paths = set()
    for batch_dir in sorted(downloads_path.iterdir()):
        if batch_dir.is_dir():
            for docx_file in batch_dir.iterdir():
                if docx_file.suffix.lower() == '.docx' and docx_file not in seen_paths:
                    docx_files.append(docx_file)
                    seen_paths.add(docx_file)

    print(f"  Found {len(docx_files)} DOCX files in downloads/")

    # Match and assemble
    print("\nMatching files to articles...")
    matched = []
    unmatched_files = []
    unmatched_articles = set(nyt_articles.keys())

    for docx_path in docx_files:
        # Try filename-based matching first
        filename_title = docx_path.stem  # filename without extension
        norm_filename = normalize_title(filename_title)

        # Also try extracting title from DOCX content
        content_title = extract_title_from_docx(docx_path)
        norm_content = normalize_title(content_title) if content_title else ""

        # Try to match
        match = None
        match_method = None

        if norm_filename in title_lookup:
            match = title_lookup[norm_filename]
            match_method = "filename"
        elif norm_content in title_lookup:
            match = title_lookup[norm_content]
            match_method = "content"

        if match:
            # Handle potential duplicates
            if isinstance(match, list):
                # Multiple articles with same title - take first unmatched
                for article_id, article in match:
                    if article_id in unmatched_articles:
                        match = (article_id, article)
                        break
                else:
                    match = match[0]  # fallback to first

            article_id, article = match
            unmatched_articles.discard(article_id)

            # Extract text
            body_text = extract_text_from_docx(docx_path)

            # Extract labels
            frame_labels = extract_frame_labels(article)
            tone_labels = extract_tone_labels(article)

            matched.append({
                "article_id": article_id,
                "title": article.get("title"),
                "year": article.get("year"),
                "month": article.get("month"),
                "source": article.get("source"),
                "byline": article.get("byline"),
                "section": article.get("section"),
                "length": article.get("length"),
                "text": body_text,
                "text_length": len(body_text),
                "docx_file": str(docx_path.relative_to(base_path)),
                "match_method": match_method,
                "frame_annotations": json.dumps(frame_labels),
                "tone_annotations": json.dumps(tone_labels),
            })
        else:
            unmatched_files.append({
                "file": str(docx_path.relative_to(base_path)),
                "filename_title": filename_title,
                "content_title": content_title
            })

    print(f"  Matched: {len(matched)}")
    print(f"  Unmatched files: {len(unmatched_files)}")
    print(f"  Missing articles: {len(unmatched_articles)}")

    # Build DataFrame
    if matched:
        df = pd.DataFrame(matched)

        # Fix column types
        df['length'] = pd.to_numeric(df['length'], errors='coerce')
        df['year'] = pd.to_numeric(df['year'], errors='coerce').astype('Int64')
        df['month'] = pd.to_numeric(df['month'], errors='coerce').astype('Int64')

        # Save outputs
        df.to_parquet(config["output_parquet"], index=False)
        df.to_csv(config["output_csv"], index=False)
        print(f"\nSaved {len(df)} articles to:")
        print(f"  - {config['output_parquet']}")
        print(f"  - {config['output_csv']}")

    # Save report
    report = {
        "total_nyt_articles": len(nyt_articles),
        "docx_files_found": len(docx_files),
        "matched": len(matched),
        "unmatched_files": unmatched_files[:50],  # truncate for readability
        "unmatched_files_count": len(unmatched_files),
        "missing_articles_count": len(unmatched_articles),
        "missing_articles_sample": list(unmatched_articles)[:20],
    }

    with open(config["report_file"], 'w', encoding='utf-8') as f:
        json.dump(report, f, indent=2)
    print(f"  - {config['report_file']}")


if __name__ == "__main__":
    main()
